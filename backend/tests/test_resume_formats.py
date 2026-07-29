"""
PDF and DOCX parsing were part of the original spec ("resumes can be
PDF/DOCX/TXT") but only TXT was ever actually tested through Phase 4.
These tests generate real PDF/DOCX files and verify extraction works —
not just that the code path exists.
"""
import io

import docx as docx_lib
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def _build_test_pdf() -> bytes:
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    lines = [
        "Ravi Kumar",
        "ravi.kumar@example.com",
        "9123456780",
        "Skills: Python, SQL, AWS, Docker",
        "Education: B.Tech Computer Science, PQR University",
    ]
    y = 750
    for line in lines:
        c.drawString(50, y, line)
        y -= 20
    c.save()
    buffer.seek(0)
    return buffer.read()


def _build_test_docx() -> bytes:
    buffer = io.BytesIO()
    document = docx_lib.Document()
    document.add_paragraph("Meera Nair")
    document.add_paragraph("meera.nair@example.com")
    document.add_paragraph("9988776655")
    document.add_paragraph("Skills: Java, Spring Boot, Docker, System Design")
    document.add_paragraph("Education: M.Tech, LMN University")
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _auth_headers(client, email):
    client.post(
        "/api/v1/auth/register",
        json={"email": email, "password": "SecurePass123", "full_name": "Format Test"},
    )
    login_response = client.post(
        "/api/v1/auth/login", data={"username": email, "password": "SecurePass123"}
    )
    token = login_response.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


def test_upload_and_parse_real_pdf(client):
    headers = _auth_headers(client, "pdfresume@example.com")
    pdf_bytes = _build_test_pdf()

    response = client.post(
        "/api/v1/resumes/upload",
        headers=headers,
        files={"file": ("resume.pdf", pdf_bytes, "application/pdf")},
    )
    assert response.status_code == 201, response.text
    data = response.json()["extracted_data"]
    assert data["email"] == "ravi.kumar@example.com"
    assert "python" in data["skills"]
    assert "aws" in data["skills"]


def test_upload_and_parse_real_docx(client):
    headers = _auth_headers(client, "docxresume@example.com")
    docx_bytes = _build_test_docx()

    response = client.post(
        "/api/v1/resumes/upload",
        headers=headers,
        files={
            "file": (
                "resume.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 201, response.text
    data = response.json()["extracted_data"]
    assert data["email"] == "meera.nair@example.com"
    assert "java" in data["skills"]
    assert "docker" in data["skills"]
    assert "system design" in data["skills"]
